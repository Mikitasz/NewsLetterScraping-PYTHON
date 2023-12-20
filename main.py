import requests
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import RGBColor, Pt
import os

from docx.shared import Inches
from docx import Document
from bs4 import BeautifulSoup 

image_folder = "images"
document = Document()
#url= "https://thehackernews.com/2023/12/remote-encryption-attacks-surge-how-one.html"
#headers={
#"Accept":"*/*",
#    "User-Agent":"Mozilla/5.0 (Windows NT 10.0; Win64; x64; rv:120.0) Gecko/20100101 Firefox/120.0",
#}
#req = requests.get(url)
#src=req.text


#with open("index.html", "w") as file:
#    file.write(src)


with open("index.html") as file:
    src=file.read()


soup= BeautifulSoup(src,'lxml')



title=soup.find_all(class_="story-title")
for item in title:
    item_text=item.text
    item_link=item.a.get("href")
title = document.add_paragraph(item_text)


text=soup.find_all(class_="articlebody")
for item in text:
    item_main_text=item.text


text = document.add_paragraph(item_main_text)




images=soup.find_all(class_="separator")


for img_tag in images[1:]:

    img_src = img_tag.find('img')
    img_link = img_src.get('data-src')
 
    response = requests.get(img_link)
    img_filename = os.path.join(image_folder,img_link.split("/")[-1])
    with open(img_filename, 'wb') as f:
        f.write(response.content)




link = document.add_paragraph("Link "+item_link)



run = title.runs[0]
font = run.font
font.color.rgb = RGBColor(0, 0, 0)  # RGB values for black
title.alignment = 0  # 0 for left-align 
font.size = Pt(16)  # Font size 16 pointss

document.save('example_document.docx')