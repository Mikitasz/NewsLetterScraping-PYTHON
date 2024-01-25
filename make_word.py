from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import RGBColor, Pt
from docx.shared import Inches
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
import docx
import os
from docxcompose.composer import Composer
from datetime import date

num_files = 0
#init date

todays_date = date.today() 


class WordDocx:
    def __init__(self, images_folder, item_link, titlestext, item_main_text, tags, li_text_polish, imgname, n,
                 h2_polish) -> None:
        self._images_folder = images_folder
        self._li_text_polish = li_text_polish
        self._tags = tags
        self._item_link = item_link
        self.__count = 0
        self._titlestext = titlestext
        self._item_main_text = item_main_text
        self._imgname = imgname
        self.main_text_font_size = 12
        self.__n = n
 
        self.__h2_polish = h2_polish

    def add_hyperlink(self, paragraph, text, url):
        print("-- Making hyperlink")
        part = paragraph.part
        r_id = part.relate_to(url, docx.opc.constants.RELATIONSHIP_TYPE.HYPERLINK, is_external=True)

        hyperlink = docx.oxml.shared.OxmlElement('w:hyperlink')
        hyperlink.set(docx.oxml.shared.qn('r:id'), r_id, )

        new_run = docx.text.run.Run(
            docx.oxml.shared.OxmlElement('w:r'), paragraph)
        
        new_run.font.color.rgb = docx.shared.RGBColor(0, 0, 255)
        new_run.text = text

        hyperlink.append(new_run._element)
        paragraph._p.append(hyperlink)
        return hyperlink

    def word_format(self):

        document = Document()
        title = document.add_heading(f"{self.__n + 1}. " + str(self._titlestext), level=1)
        j = 0
        img_count = 0
        img_index = 0
        print("-- Creating word-docx file")
        for i in range(len(self._tags)):

            if self._tags[i] == "p":
                text = document.add_paragraph(str(self._item_main_text[j]))
                j += 1
                run_main = text.runs[0]
                font_main = run_main.font
                font_main.color.rgb = RGBColor(0, 0, 0)
                font_main.name = 'Calibri'
                font_main = Pt(11)  # F
            if self._tags[i] == "li":
                li = document.add_paragraph(str(self._li_text_polish[self.__count])[33:-49])
                self.__count += 1
                li.paragraph_format.left_indent = Pt(30)
                run_main = li.runs[0]
                font_main = run_main.font
                font_main.color.rgb = RGBColor(0, 0, 0)
                font_main.name = 'Calibri'
                font_main = Pt(11)  # F
            if self._tags[i] == "h2":
                text = document.add_paragraph(str(self.__h2_polish)[33:-49])
                text.paragraph_format.left_indent = Pt(30)
                run_main = text.runs[0]
                font_main = run_main.font
                font_main.bold = True
                font_main.color.rgb = RGBColor(0, 0, 0)
                font_main.name = 'Calibri'
                font_main = Pt(13)  # F
            if self._tags[i] == "img" and img_count != 0:
                document.add_picture(self._imgname[img_index], width=Inches(4), height=Inches(3))
                last_paragraph = document.paragraphs[-1]
                last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                img_count += 1
                img_index += 1
            else:
                img_count += 1
        link = document.add_paragraph("Link: ")


        self.add_hyperlink(link, self._item_link, self._item_link)
        run = title.runs[0]

        font = run.font
        font.color.rgb = RGBColor(0, 0, 0)  # RGB values for black
        title.alignment = 0  # 0 for left-align 
        font.name = 'Calibri'
        font.size = Pt(self.main_text_font_size)  # Font size 16 pointss
        global num_files

        num_files += 1
        document.save(f"{self.__n}file.docx")

    def merge(self):
        global num_files

        master = Document("0file.docx")
        composer = Composer(master)
        for i in range(1, num_files):
            doc_temp = Document(f"{i}file.docx")
            composer.append(doc_temp)
        composer.save(f"Final file is Newsletter {todays_date.day}.{todays_date.month}.docx")

        for i in range(0, num_files):
            os.remove(f"{i}file.docx")
